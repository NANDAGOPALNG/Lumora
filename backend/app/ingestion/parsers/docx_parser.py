"""DOCX paragraph extraction.

Extracts paragraphs in document order; embedded images are ignored.
DOCX has no reliable page boundaries without rendering the layout, so
the whole document is treated as a single page.
"""

from docx import Document as DocxDocument

from app.ingestion.parsers.types import ParsedDocument, ParsedPage, ParsingError


def parse(file_path: str) -> ParsedDocument:
    try:
        docx_document = DocxDocument(file_path)
    except Exception as exc:
        raise ParsingError(f"Could not read DOCX file: {exc}") from exc

    paragraphs = [p.text for p in docx_document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    if not text.strip():
        raise ParsingError("DOCX contains no extractable text")

    return ParsedDocument(pages=[ParsedPage(text=text, page_number=None)])
